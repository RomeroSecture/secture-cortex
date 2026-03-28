"""Text extraction from binary file formats (PDF, DOCX).

Sync functions designed to run via asyncio.to_thread() to avoid
blocking the event loop. Each returns extracted text or empty string
on failure — never raises.
"""

import contextlib
import io
import re

import structlog

logger = structlog.get_logger()


def extract_pdf_text(content: bytes) -> str:
    """Extract text from a PDF file.

    Uses PyMuPDF for fast, accurate extraction with table support.
    Handles multi-column layouts via sort=True.
    Returns empty string on any failure.
    """
    try:
        import pymupdf

        doc = pymupdf.open(stream=content, filetype="pdf")

        if not doc.is_pdf:
            logger.warning("pdf_not_valid", msg="Content is not a valid PDF")
            doc.close()
            return ""

        pages_text: list[str] = []
        tables_found = 0

        for _page_num, page in enumerate(doc):
            # Extract text with layout-aware sorting
            text = page.get_text(sort=True)

            # Extract tables if available
            try:
                tables = page.find_tables()
                for table in tables:
                    table_data = table.extract()
                    if table_data:
                        tables_found += 1
                        rows = []
                        for row in table_data:
                            cells = [
                                str(c).strip() if c else ""
                                for c in row
                            ]
                            rows.append(" | ".join(cells))
                        text += "\n\n" + "\n".join(rows)
            except Exception:
                pass  # Tables API may not be available on all pages

            if text.strip():
                pages_text.append(text.strip())

        doc.close()

        full_text = "\n\n".join(pages_text)
        # Collapse excessive whitespace
        full_text = re.sub(r"\n{3,}", "\n\n", full_text)
        full_text = full_text.strip()

        word_count = len(full_text.split()) if full_text else 0
        logger.info(
            "pdf_extracted",
            pages=len(pages_text),
            words=word_count,
            tables=tables_found,
        )
        return full_text

    except Exception:
        logger.exception("pdf_extraction_failed")
        return ""


def extract_docx_text(content: bytes) -> str:
    """Extract text from a DOCX file.

    Uses python-docx with body element iteration to preserve
    paragraph/table ordering. Headings get ## prefix for
    natural chunk boundaries.
    Returns empty string on any failure.
    """
    try:
        from docx import Document
        from docx.oxml.ns import qn

        doc = Document(io.BytesIO(content))
        parts: list[str] = []

        # Build lookup maps from properly-constructed objects
        para_map = {id(p._element): p for p in doc.paragraphs}
        table_map = {id(t._element): t for t in doc.tables}

        # Iterate body children to preserve document order
        for element in doc.element.body:
            elem_id = id(element)

            if elem_id in para_map:
                para = para_map[elem_id]
                text = para.text.strip()
                if not text:
                    continue

                # Prefix headings for natural chunk boundaries
                style_name = ""
                with contextlib.suppress(Exception):
                    style_name = para.style.name if para.style else ""
                if style_name.startswith("Heading"):
                    level = style_name.replace("Heading", "").strip()
                    prefix = "#" * (int(level) if level.isdigit() else 2)
                    parts.append(f"{prefix} {text}")
                else:
                    parts.append(text)

            elif elem_id in table_map:
                table = table_map[elem_id]
                rows: list[str] = []
                for row in table.rows:
                    cells = [
                        cell.text.strip() for cell in row.cells
                    ]
                    rows.append(" | ".join(cells))
                if rows:
                    parts.append("\n".join(rows))

            # Also check for raw w:p not in paragraphs list
            elif element.tag == qn("w:p"):
                text = element.text or ""
                for child in element.iter():
                    if child.text:
                        text += child.text
                    if child.tail:
                        text += child.tail
                text = text.strip()
                if text:
                    parts.append(text)

        full_text = "\n\n".join(parts)
        full_text = re.sub(r"\n{3,}", "\n\n", full_text)
        full_text = full_text.strip()

        word_count = len(full_text.split()) if full_text else 0
        para_count = sum(
            1 for p in doc.paragraphs if p.text.strip()
        )
        table_count = len(doc.tables)
        logger.info(
            "docx_extracted",
            paragraphs=para_count,
            tables=table_count,
            words=word_count,
        )
        return full_text

    except Exception:
        logger.exception("docx_extraction_failed")
        return ""
