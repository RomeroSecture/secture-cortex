"""Tests for PDF and DOCX text extraction (RAG pipeline fix)."""

import io

import pytest

from app.services.extraction import extract_docx_text, extract_pdf_text

# ─── Helper: generate test PDFs ───────────────────────────────


def _make_pdf(text: str, pages: int = 1) -> bytes:
    """Generate a minimal PDF with given text using fpdf2."""
    from fpdf import FPDF

    pdf = FPDF()
    for _ in range(pages):
        pdf.add_page()
        pdf.set_font("Helvetica", size=12)
        pdf.multi_cell(0, 10, text)
    return bytes(pdf.output())


def _make_docx(
    paragraphs: list[str],
    heading: str = "",
    table_data: list[list[str]] | None = None,
) -> bytes:
    """Generate a minimal DOCX with paragraphs, optional heading and table."""
    from docx import Document

    doc = Document()
    if heading:
        doc.add_heading(heading, level=1)
    for p in paragraphs:
        doc.add_paragraph(p)
    if table_data:
        rows = len(table_data)
        cols = len(table_data[0]) if table_data else 0
        table = doc.add_table(rows=rows, cols=cols)
        for i, row_data in enumerate(table_data):
            for j, cell_text in enumerate(row_data):
                table.rows[i].cells[j].text = cell_text
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


# ─── PDF extraction tests ─────────────────────────────────────


def test_extract_pdf_basic() -> None:
    """Single-page PDF with text extracts correctly."""
    pdf_bytes = _make_pdf("Este es un documento de prueba para RAG")
    text = extract_pdf_text(pdf_bytes)
    assert "documento" in text.lower()
    assert "prueba" in text.lower()
    assert len(text.split()) >= 5


def test_extract_pdf_multi_page() -> None:
    """Multi-page PDF extracts text from all pages."""
    pdf_bytes = _make_pdf("Página con contenido importante", pages=3)
    text = extract_pdf_text(pdf_bytes)
    # Should contain content from multiple pages
    assert text.count("contenido") >= 2


def test_extract_pdf_empty_page() -> None:
    """PDF with no text returns empty or whitespace-only."""
    from fpdf import FPDF

    pdf = FPDF()
    pdf.add_page()  # Empty page, no text
    pdf_bytes = bytes(pdf.output())
    text = extract_pdf_text(pdf_bytes)
    assert text.strip() == "" or len(text.split()) < 3


def test_extract_pdf_invalid_bytes() -> None:
    """Random bytes → empty string, no crash."""
    text = extract_pdf_text(b"this is not a pdf at all")
    assert text == ""


def test_extract_pdf_spanish_unicode() -> None:
    """Spanish text with accents and ñ extracts correctly."""
    pdf_bytes = _make_pdf("Señor García pidió información técnica")
    text = extract_pdf_text(pdf_bytes)
    assert "García" in text or "garcia" in text.lower()


# ─── DOCX extraction tests ────────────────────────────────────


def test_extract_docx_basic() -> None:
    """DOCX with paragraphs extracts all text."""
    docx_bytes = _make_docx([
        "Primer párrafo del documento.",
        "Segundo párrafo con más información.",
        "Tercer párrafo sobre el proyecto.",
    ])
    text = extract_docx_text(docx_bytes)
    assert "Primer" in text
    assert "Segundo" in text
    assert "Tercer" in text


def test_extract_docx_with_heading() -> None:
    """DOCX heading gets ## prefix for chunk boundaries."""
    docx_bytes = _make_docx(
        ["Contenido bajo el heading."],
        heading="Arquitectura del Sistema",
    )
    text = extract_docx_text(docx_bytes)
    assert "Arquitectura" in text
    assert "#" in text  # Heading should have markdown prefix


def test_extract_docx_with_table() -> None:
    """DOCX table content appears in extracted text."""
    docx_bytes = _make_docx(
        ["Datos del proyecto:"],
        table_data=[
            ["Componente", "Estado"],
            ["Backend", "Completo"],
            ["Frontend", "En progreso"],
        ],
    )
    text = extract_docx_text(docx_bytes)
    assert "Backend" in text
    assert "Frontend" in text
    assert "Completo" in text


def test_extract_docx_invalid_bytes() -> None:
    """Random bytes → empty string, no crash."""
    text = extract_docx_text(b"not a docx file")
    assert text == ""


# ─── Binary validation tests ──────────────────────────────────


def test_validate_binary_pdf_magic_bytes() -> None:
    """Real PDF starts with %PDF- magic bytes."""
    pdf_bytes = _make_pdf("test")
    assert pdf_bytes[:5] == b"%PDF-"


def test_validate_binary_docx_magic_bytes() -> None:
    """Real DOCX starts with PK zip magic bytes."""
    docx_bytes = _make_docx(["test"])
    assert docx_bytes[:4] == b"PK\x03\x04"


@pytest.mark.asyncio
async def test_validate_binary_content_rejects_fake_pdf() -> None:
    """_validate_binary_content raises 422 for fake PDF."""
    from fastapi import HTTPException

    from app.api.v1.context import _validate_binary_content

    with pytest.raises(HTTPException) as exc_info:
        _validate_binary_content(b"JPEG fake content", "test.pdf")
    assert exc_info.value.status_code == 422


@pytest.mark.asyncio
async def test_extract_text_async_pdf() -> None:
    """_extract_text async wrapper works for PDF."""
    from app.api.v1.context import _extract_text

    pdf_bytes = _make_pdf("Texto de prueba async")
    text = await _extract_text(pdf_bytes, "test.pdf")
    assert "prueba" in text.lower()


@pytest.mark.asyncio
async def test_extract_text_async_docx() -> None:
    """_extract_text async wrapper works for DOCX."""
    from app.api.v1.context import _extract_text

    docx_bytes = _make_docx(["Documento DOCX async test"])
    text = await _extract_text(docx_bytes, "test.docx")
    assert "DOCX" in text


@pytest.mark.asyncio
async def test_extract_text_txt_unchanged() -> None:
    """Plain text extraction still works as before."""
    from app.api.v1.context import _extract_text

    text = await _extract_text(b"Hello world", "test.txt")
    assert text == "Hello world"
